import asyncio
import os
import aiopg
import hashlib
import json
import aiofiles
import streamlit as st
from typing import List
from openai import AsyncOpenAI
from dotenv import load_dotenv
import fitz
from docx import Document
import io
import pytesseract
from PIL import Image
import re
import tiktoken
import time
import zipfile
import logging
from tenacity import retry, stop_after_attempt, wait_exponential

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
SysPromptDefault = "You are now in the role of an expert AI."

# Load prompts from file
try:
    with open('prompt.json', 'r') as file:
        prompts = json.load(file)
    resume_prompt = prompts['resume_prompt']
except FileNotFoundError:
    logger.error("prompt.json file not found.")
    resume_prompt = "Analyze the following resume and extract key information."
except json.JSONDecodeError:
    logger.error("Error decoding prompt.json. Using default prompt.")
    resume_prompt = "Analyze the following resume and extract key information."


def get_digest(content):
    h = hashlib.sha256()
    if isinstance(content, str):
        content = content.encode()
    h.update(content)
    return h.hexdigest()

async def get_db_pool():
    try:
        return await aiopg.create_pool(
            dsn="dbname='postgres' user='postgres.rrjupfnfhyikdeitbktn' password='nI20th0in3@' host='aws-0-us-east-1.pooler.supabase.com' port='6543'"
        )
    except Exception as e:
        logger.error(f"Failed to create database pool: {e}")
        raise


async def is_file_processed_async(file_id):
    try:
        async with await get_db_pool() as pool:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("SELECT 1 FROM resume_info WHERE file_id = %s", (file_id,))
                    return await cur.fetchone() is not None
    except Exception as e:
        logger.error(f"Error checking if file is processed: {e}")
        return False


def count_tokens(text: str) -> int:
    try:
        # Use cl100k_base encoding, which is suitable for many modern models
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))
    except Exception as e:
        logger.error(f"Error counting tokens: {e}")
        # Fallback to a rough estimate if tiktoken fails
        return len(text.split())


@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
async def response_async(message: str, model: str = "llama3-70b-8192", SysPrompt: str = SysPromptDefault,
                         temperature: float = 0.2) -> str:
    client = AsyncOpenAI(
        api_key=GROQ_API_KEY,
        base_url="https://api.groq.com/openai/v1",
    )

    messages = [{"role": "system", "content": SysPrompt}, {"role": "user", "content": message}]
    total_tokens = sum(count_tokens(msg["content"]) for msg in messages)

    # Adjust the token limit based on Groq's specific limitations
    if total_tokens > 5500:  # This threshold might need adjustment based on Groq's actual limits
        wait_time = 60
        logger.info(f"Approaching estimated token limit. Waiting for {wait_time} seconds...")
        await asyncio.sleep(wait_time)

    try:
        response = await client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            frequency_penalty=0.2,
        )
        return response.choices[0].message.content
    except Exception as e:
        if "rate_limit_exceeded" in str(e):
            wait_time = 60
            logger.warning(f"Rate limit exceeded. Waiting for {wait_time} seconds...")
            await asyncio.sleep(wait_time)
            raise  # Retry will be handled by the decorator
        else:
            logger.error(f"Error in API response: {e}")
            raise


async def extract_content_async(file_content: bytes, file_type: str) -> List[str]:
    def extract_text_from_image(image: Image.Image) -> str:
        try:
            return pytesseract.image_to_string(image)
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            return ""

    try:
        if file_type == 'pdf':
            pdf_doc = fitz.open(stream=file_content, filetype="pdf")
            pages_content = []
            for page in pdf_doc:
                text_content = page.get_text("text")
                if not text_content.strip():
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text_content = extract_text_from_image(img)
                pages_content.append(text_content.replace("\n", "\t"))
            pdf_doc.close()
            return pages_content

        elif file_type == 'docx':
            doc = Document(io.BytesIO(file_content))
            content = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                content.extend(['\t'.join(cell.text for cell in row.cells) for row in table.rows])
            return content

        elif file_type in ['jpeg', 'png']:
            img = Image.open(io.BytesIO(file_content))
            return [extract_text_from_image(img).replace("\n", "\t")]

        else:
            raise ValueError("Unsupported file type. Please provide a PDF, DOCX, JPEG, or PNG file.")
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        return []


async def insert_data_resume_async(file_name, file_id, resume):
    try:
        async with await get_db_pool() as pool:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    resume_json = json.dumps(resume)
                    await cur.execute("""
                        INSERT INTO resume_info(file_id, file_name, resume_parse)
                        VALUES (%s, %s, %s)
                        ON CONFLICT (file_id) DO NOTHING;
                    """, (file_id, file_name, resume_json))
    except Exception as e:
        logger.error(f"Error inserting data into database: {e}")
        raise


def extract_json(response_str):
    response_str = re.sub(r'//.*?\n|/\*.*?\*/', '', response_str, flags=re.DOTALL)
    match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", response_str, re.DOTALL) or re.search(r"\{.*\}", response_str,
                                                                                               re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON detected. Error: {e}")
    logger.warning("No valid JSON part found in the response string.")
    return None


async def identification_async(file_content, file_type):
    data = await extract_content_async(file_content, file_type)
    if not data:
        logger.warning("No content could be extracted from the file.")
        return None
    context = "\n\n".join(data)
    message = f"RESUME\n\n{context}\n\n"

    max_retries = 3
    for attempt in range(max_retries):
        try:
            response_str = await response_async(message=message, model="llama3-70b-8192", SysPrompt=resume_prompt,
                                                temperature=0)
            return extract_json(response_str)
        except Exception as e:
            if attempt < max_retries - 1:
                logger.warning(f"Error occurred: {e}. Retrying in 60 seconds...")
                await asyncio.sleep(60)
            else:
                logger.error(f"Max retries reached. Failed to process file.")
                return None


async def process_single_file(file_path):
    try:
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower()

        if file_ext not in ['.pdf', '.docx', '.jpeg', '.jpg', '.png']:
            logger.info(f"Skipping unsupported file type: {file_name}")
            return

        file_type = 'pdf' if file_ext == '.pdf' else 'docx' if file_ext == '.docx' else 'jpeg' if file_ext in ['.jpeg',
                                                                                                               '.jpg'] else 'png'

        async with aiofiles.open(file_path, mode='rb') as f:
            file_content = await f.read()

        file_id = get_digest(file_content)
        if await is_file_processed_async(file_id):
            logger.info(f"File already processed: {file_name}")
            return

        final_output = await identification_async(file_content, file_type)
        if final_output:
            await insert_data_resume_async(file_name, file_id, final_output)
            logger.info(f"Successfully processed: {file_name}")
        else:
            logger.warning(f"Failed to process: {file_name}")
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {e}")


async def process_batch(batch_files):
    tasks = [process_single_file(file) for file in batch_files]
    await asyncio.gather(*tasks)


async def process_files_in_directory_async(directory):
    files = []
    for root, _, filenames in os.walk(directory):
        for filename in filenames:
            files.append(os.path.join(root, filename))

    num_files = len(files)
    if num_files == 0:
        st.warning("No files found in the selected directory.")
        return

    batch_size = 100  # Process 100 files at once
    batches = [files[i:i + batch_size] for i in range(0, len(files), batch_size)]

    for i, batch in enumerate(batches, 1):
        await process_batch(batch)
        st.write(f"Processed batch {i}/{len(batches)}")

    st.success("Processing and uploading completed successfully.")


def main():
    st.title("Resume Batch Processing")

    uploaded_file = st.file_uploader("Upload a folder (as a ZIP file):", type=["zip"])

    if uploaded_file is not None:
        try:
            # Save the uploaded ZIP file temporarily
            with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                zip_path = tmp_file.name

            # Extract the ZIP file contents
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                extract_dir = tempfile.mkdtemp()
                zip_ref.extractall(extract_dir)

            st.write("Files extracted successfully. Processing...")

            asyncio.run(process_files_in_directory_async(extract_dir))

        except Exception as e:
            st.error(f"An error occurred during processing: {e}")
            logger.exception("An error occurred during processing")
    else:
        st.warning("Please upload a ZIP file containing the resumes.")

if __name__ == "__main__":
    main()
